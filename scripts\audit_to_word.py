"""
audit_to_word.py
----------------
游戏数值策划 · 任务活动奖励审核工具

功能：
  1. 读取 xlsx 文件的所有工作表
  2. 对每张工作表执行奖励逻辑审核（由 AI 分析结果驱动）
  3. 将有问题的区域高亮截图（使用 matplotlib 渲染表格）
  4. 将「截图 + 文字描述」整合进 Word 审核报告

用法：
  python audit_to_word.py <xlsx路径> [输出目录] [--issues JSON字符串]

  --issues 参数格式（由 AI 分析后传入）：
  [
    {
      "sheet": "Sheet1",
      "row": 3,                  # 1-based，含表头，-1 表示整列/整体问题
      "col": 2,                  # 1-based，-1 表示整行/整体问题
      "rows": [3, 4],            # 可选，多行范围 [start, end]（含）
      "cols": [1, 5],            # 可选，多列范围 [start, end]（含）
      "severity": "严重",        # 严重 / 中等 / 建议
      "title": "计算错误",
      "description": "第3行总奖励(500)与子项加总(300+150=450)不符",
      "suggestion": "将总奖励修改为450"
    },
    ...
  ]

依赖：
  pip install openpyxl pandas matplotlib pillow python-docx
"""

import sys
import json
import argparse
import os
from datetime import datetime
from pathlib import Path

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib import rcParams
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 字体配置（支持中文）────────────────────────────────────────────
def _setup_font():
    """尝试设置支持中文的字体"""
    candidates = [
        "Microsoft YaHei", "SimHei", "SimSun",
        "PingFang SC", "Hiragino Sans GB", "WenQuanYi Micro Hei",
        "DejaVu Sans",
    ]
    for font in candidates:
        try:
            rcParams["font.family"] = font
            rcParams["axes.unicode_minus"] = False
            return font
        except Exception:
            continue
    return "DejaVu Sans"


FONT_NAME = _setup_font()

# ── 严重等级配置 ───────────────────────────────────────────────────
SEVERITY_CONFIG = {
    "严重": {"emoji": "🔴", "color": "#FF4444", "bg": "#FFE5E5"},
    "中等": {"emoji": "🟡", "color": "#FFA500", "bg": "#FFF8E1"},
    "建议": {"emoji": "🟢", "color": "#4CAF50", "bg": "#E8F5E9"},
}


# ══════════════════════════════════════════════════════════════════
#  截图生成
# ══════════════════════════════════════════════════════════════════

def _col_letter(n: int) -> str:
    """将 1-based 列号转为字母（1→A, 27→AA）"""
    result = ""
    while n > 0:
        n, r = divmod(n - 1, 26)
        result = chr(65 + r) + result
    return result


def render_sheet_screenshot(
    df: pd.DataFrame,
    sheet_name: str,
    issue: dict,
    output_path: str,
    context_rows: int = 3,
) -> str:
    """
    将 DataFrame 的问题区域渲染为带高亮的表格截图。

    参数：
        df            - 该工作表的 DataFrame（含表头，index 从 0 开始）
        sheet_name    - 工作表名称
        issue         - 单条问题描述 dict
        output_path   - 截图保存路径（.png）
        context_rows  - 问题行上下各展示多少行上下文

    返回：
        截图文件路径
    """
    total_rows, total_cols = df.shape

    # ── 确定高亮行列范围（1-based 转 0-based）──────────────────────
    if "rows" in issue:
        hi_row_start = max(0, issue["rows"][0] - 2)   # -1 for header, -1 for 0-based
        hi_row_end   = min(total_rows - 1, issue["rows"][1] - 2)
    elif issue.get("row", -1) != -1:
        hi_row_start = hi_row_end = issue["row"] - 2
    else:
        hi_row_start = hi_row_end = -1  # 整体问题，不高亮特定行

    if "cols" in issue:
        hi_col_start = issue["cols"][0] - 1
        hi_col_end   = issue["cols"][1] - 1
    elif issue.get("col", -1) != -1:
        hi_col_start = hi_col_end = issue["col"] - 1
    else:
        hi_col_start = hi_col_end = -1

    # ── 确定展示的行范围 ────────────────────────────────────────────
    if hi_row_start >= 0:
        view_start = max(0, hi_row_start - context_rows)
        view_end   = min(total_rows - 1, hi_row_end + context_rows)
    else:
        view_start = 0
        view_end   = min(total_rows - 1, 20)

    sub_df = df.iloc[view_start : view_end + 1].reset_index(drop=True)
    n_rows = len(sub_df)
    n_cols = len(sub_df.columns)

    # ── 截断过长的单元格内容 ─────────────────────────────────────────
    MAX_CELL_LEN = 20
    display_df = sub_df.copy()
    for c in display_df.columns:
        display_df[c] = display_df[c].astype(str).apply(
            lambda x: x[:MAX_CELL_LEN] + "…" if len(x) > MAX_CELL_LEN else x
        )

    # ── 计算图像尺寸 ─────────────────────────────────────────────────
    cell_w = max(1.2, 14 / n_cols)
    cell_h = 0.45
    fig_w  = n_cols * cell_w + 0.5
    fig_h  = (n_rows + 1) * cell_h + 1.0   # +1 for header

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")

    severity = issue.get("severity", "中等")
    cfg = SEVERITY_CONFIG.get(severity, SEVERITY_CONFIG["中等"])

    # ── 绘制表格 ─────────────────────────────────────────────────────
    col_labels = [str(c) for c in display_df.columns]
    cell_text  = display_df.values.tolist()

    tbl = ax.table(
        cellText=cell_text,
        colLabels=col_labels,
        cellLoc="center",
        loc="center",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)
    tbl.scale(1, 1.4)

    # 表头样式
    for j in range(n_cols):
        cell = tbl[0, j]
        cell.set_facecolor("#2C3E50")
        cell.get_text().set_color("white")
        cell.get_text().set_fontweight("bold")

    # 高亮问题单元格
    for i in range(n_rows):
        abs_row = view_start + i
        for j in range(n_cols):
            cell = tbl[i + 1, j]
            cell.set_facecolor("#F8F9FA")

            if hi_row_start >= 0 and hi_row_end >= 0:
                row_hit = hi_row_start <= abs_row <= hi_row_end
            else:
                row_hit = True

            if hi_col_start >= 0 and hi_col_end >= 0:
                col_hit = hi_col_start <= j <= hi_col_end
            else:
                col_hit = True

            if row_hit and col_hit:
                cell.set_facecolor(cfg["bg"])
                cell.set_edgecolor(cfg["color"])
                cell.set_linewidth(2)

    # 行号标注
    for i in range(n_rows):
        abs_row = view_start + i
        tbl[i + 1, 0].get_text().set_text(
            f"[{abs_row + 2}] {str(display_df.iloc[i, 0])}"
            if n_cols > 0 else f"[{abs_row + 2}]"
        )

    # 标题
    title_text = (
        f"{cfg['emoji']} [{severity}] {issue.get('title', '')}  "
        f"· 工作表：{sheet_name}"
    )
    ax.set_title(title_text, fontsize=10, fontweight="bold",
                 color=cfg["color"], pad=12)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    return output_path


# ══════════════════════════════════════════════════════════════════
#  Word 报告生成
# ══════════════════════════════════════════════════════════════════

def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_colored_paragraph(doc: Document, text: str, color_hex: str, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor(
        int(color_hex[1:3], 16),
        int(color_hex[3:5], 16),
        int(color_hex[5:7], 16),
    )
    return p


def _set_table_style(table):
    """给 Word 表格加基本边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "AAAAAA")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def generate_word_report(
    xlsx_path: str,
    issues: list,
    output_dir: str,
) -> str:
    """
    生成带截图的 Word 审核报告。

    参数：
        xlsx_path  - 原始 xlsx 文件路径
        issues     - 问题列表（由 AI 分析后传入）
        output_dir - 输出目录

    返回：
        生成的 Word 文件路径
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    screenshots_dir = output_dir / "screenshots"
    screenshots_dir.mkdir(exist_ok=True)

    # 读取所有工作表
    xlsx_path = Path(xlsx_path)
    xls = pd.ExcelFile(xlsx_path)
    sheet_names = xls.sheet_names
    sheets_data = {name: xls.parse(name, header=0) for name in sheet_names}

    # 统计
    total_issues = len(issues)
    severity_count = {"严重": 0, "中等": 0, "建议": 0}
    for iss in issues:
        sev = iss.get("severity", "中等")
        severity_count[sev] = severity_count.get(sev, 0) + 1

    # ── 创建 Word 文档 ────────────────────────────────────────────
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width  = int(11906)   # A4 宽 (twips)
    section.page_height = int(16838)   # A4 高 (twips)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # 封面标题
    title_p = doc.add_heading("奖励配置审核报告", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"文件：{xlsx_path.name}").font.size = Pt(11)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.add_run(f"审核时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(10)

    doc.add_paragraph()

    # 总览表格
    _add_heading(doc, "一、总体概况", 1)
    overview_table = doc.add_table(rows=2, cols=5)
    _set_table_style(overview_table)
    headers = ["审核工作表数", "问题总数", "🔴 严重", "🟡 中等", "🟢 建议"]
    values  = [
        str(len(sheet_names)),
        str(total_issues),
        str(severity_count["严重"]),
        str(severity_count["中等"]),
        str(severity_count["建议"]),
    ]
    for j, h in enumerate(headers):
        cell = overview_table.cell(0, j)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for j, v in enumerate(values):
        overview_table.cell(1, j).text = v

    doc.add_paragraph()

    # 逐工作表问题
    _add_heading(doc, "二、逐工作表审核结果", 1)

    # 按 sheet 分组
    issues_by_sheet: dict = {}
    for iss in issues:
        s = iss.get("sheet", "未知工作表")
        issues_by_sheet.setdefault(s, []).append(iss)

    issue_global_idx = 0
    for sheet_name in sheet_names:
        sheet_issues = issues_by_sheet.get(sheet_name, [])

        _add_heading(doc, f"【{sheet_name}】", 2)

        if not sheet_issues:
            doc.add_paragraph("✅ 该工作表未发现明显问题。")
            doc.add_paragraph()
            continue

        df = sheets_data.get(sheet_name, pd.DataFrame())

        for idx, iss in enumerate(sheet_issues):
            issue_global_idx += 1
            severity = iss.get("severity", "中等")
            cfg = SEVERITY_CONFIG.get(severity, SEVERITY_CONFIG["中等"])

            # 问题标题行
            title_str = (
                f"问题 {issue_global_idx}：{cfg['emoji']} [{severity}]  "
                f"{iss.get('title', '（未命名问题）')}"
            )
            _add_colored_paragraph(doc, title_str, cfg["color"], bold=True)

            # 截图
            if not df.empty:
                shot_path = str(
                    screenshots_dir / f"issue_{issue_global_idx:03d}_{sheet_name}.png"
                )
                try:
                    render_sheet_screenshot(df, sheet_name, iss, shot_path)
                    doc.add_picture(shot_path, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"（截图生成失败：{e}）")
            else:
                doc.add_paragraph("（工作表数据为空，无法生成截图）")

            # 文字描述
            desc_table = doc.add_table(rows=3, cols=2)
            _set_table_style(desc_table)
            rows_data = [
                ("错误描述", iss.get("description", "")),
                ("影响分析", iss.get("impact", "待补充")),
                ("修改建议", iss.get("suggestion", "")),
            ]
            for r_idx, (label, content) in enumerate(rows_data):
                label_cell = desc_table.cell(r_idx, 0)
                label_cell.text = label
                label_cell.paragraphs[0].runs[0].bold = True
                label_cell.width = Inches(1.2)
                desc_table.cell(r_idx, 1).text = content

            doc.add_paragraph()

    # 汇总建议
    _add_heading(doc, "三、汇总建议", 1)
    if total_issues == 0:
        doc.add_paragraph("整体奖励配置未发现明显问题，逻辑合理。")
    else:
        doc.add_paragraph(
            f"本次审核共发现 {total_issues} 个问题，"
            f"其中严重问题 {severity_count['严重']} 个，"
            f"请优先修复严重级别问题，避免影响游戏经济平衡。"
        )
        doc.add_paragraph(
            "建议在修复后重新提交审核，确认所有计算公式、奖励上限及梯度配置均符合设计预期。"
        )

    # 保存
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"奖励审核报告_{xlsx_path.stem}_{ts}.docx"
    out_path = output_dir / out_name
    doc.save(str(out_path))
    print(f"[OK] 报告已生成：{out_path}")
    return str(out_path)


# ══════════════════════════════════════════════════════════════════
#  CLI 入口
# ══════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="游戏奖励审核 → Word 报告生成工具"
    )
    parser.add_argument("xlsx", help="xlsx 文件路径")
    parser.add_argument(
        "output_dir", nargs="?", default=None,
        help="输出目录（默认与 xlsx 同目录）"
    )
    parser.add_argument(
        "--issues", default="[]",
        help="问题列表 JSON 字符串（由 AI 分析后传入）"
    )
    args = parser.parse_args()

    xlsx_path = Path(args.xlsx).resolve()
    if not xlsx_path.exists():
        print(f"❌ 文件不存在：{xlsx_path}", file=sys.stderr)
        sys.exit(1)

    output_dir = args.output_dir or str(xlsx_path.parent)

    try:
        issues = json.loads(args.issues)
    except json.JSONDecodeError as e:
        print(f"❌ --issues 参数 JSON 解析失败：{e}", file=sys.stderr)
        sys.exit(1)

    generate_word_report(str(xlsx_path), issues, output_dir)


if __name__ == "__main__":
    main()
